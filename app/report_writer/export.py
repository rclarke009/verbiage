"""Export claim draft to DOCX (Phase 3)."""

from __future__ import annotations

import io

from docx import Document

from app.report_writer.constants import REPORT_SECTIONS


def draft_to_docx_bytes(sections: dict[str, dict], title: str = "Engineering Report") -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    for key, label in REPORT_SECTIONS:
        sec = sections.get(key) or {}
        content = (sec.get("content") or "").strip()
        if not content:
            continue
        doc.add_heading(label, level=1)
        doc.add_paragraph(content)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
